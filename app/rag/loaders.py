"""
Document loaders for the claims knowledge base.
Handles loading and parsing of PDF, DOCX, Markdown, and JSON files.
"""

import csv
import json
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional

import yaml
from pypdf import PdfReader
from docx import Document as DocxDocument
import docx2txt

from app.config import get_settings

# Knowledge base directory
KNOWLEDGE_BASE_DIR = Path(__file__).parent.parent.parent / "data" / "knowledge_base"


class Document:
    """Represents a loaded document with text and metadata."""

    def __init__(
        self,
        text: str,
        source_id: str,
        source_path: str,
        doc_type: str,
        insurance_type: str,
        product_code: Optional[str] = None,
        product_name: Optional[str] = None,
        claim_type: Optional[str] = None,
        raw_metadata: Optional[Dict[str, Any]] = None,
    ):
        self.text = text
        self.source_id = source_id
        self.source_path = source_path
        self.doc_type = doc_type
        self.insurance_type = insurance_type
        self.product_code = product_code
        self.product_name = product_name
        self.claim_type = claim_type
        self.raw_metadata = raw_metadata or {}

    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "source_id": self.source_id,
            "source_path": self.source_path,
            "doc_type": self.doc_type,
            "insurance_type": self.insurance_type,
            "product_code": self.product_code,
            "product_name": self.product_name,
            "claim_type": self.claim_type,
            "raw_metadata": self.raw_metadata,
        }


def load_manifest() -> Dict[str, Any]:
    """
    Reads manifest.yaml from KNOWLEDGE_BASE_DIR and returns the parsed dict.

    Returns:
        Parsed manifest dictionary with sources list.
    """
    manifest_path = KNOWLEDGE_BASE_DIR / "manifest.yaml"
    with open(manifest_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def iter_manifest_sources() -> Iterator[Dict[str, Any]]:
    """
    Yields a normalized structure for each source in the manifest.

    Yields:
        Dictionary with: id, full_path, doc_type, insurance_type, product_code,
        product_name, claim_type, jurisdiction, and any extra metadata.
    """
    manifest = load_manifest()
    for source in manifest.get("sources", []):
        yield {
            "id": source.get("id"),
            "full_path": str(KNOWLEDGE_BASE_DIR / source.get("path")),
            "doc_type": source.get("doc_type"),
            "insurance_type": source.get("insurance_type"),
            "product_code": source.get("product_code"),
            "product_name": source.get("product_name"),
            "claim_type": source.get("claim_type"),
            "jurisdiction": source.get("jurisdiction"),
            "metadata": {k: v for k, v in source.items() if k not in [
                "id", "path", "doc_type", "insurance_type",
                "product_code", "product_name", "claim_type", "jurisdiction"
            ]},
        }


def _load_pdf(file_path: str) -> str:
    """
    Load text from a PDF file using pypdf.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Extracted text content.
    """
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return "\n\n".join(text_parts)


def _load_docx(file_path: str) -> str:
    """
    Load text from a DOCX file using python-docx and docx2txt.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text content.
    """
    # Try docx2txt first for better text extraction
    try:
        text = docx2txt.process(file_path)
        if text and text.strip():
            return text
    except Exception:
        pass

    # Fallback to python-docx
    doc = DocxDocument(file_path)
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    return "\n\n".join(text_parts)


def _load_markdown(file_path: str) -> str:
    """
    Load text from a Markdown file via simple file read.

    Args:
        file_path: Path to the Markdown file.

    Returns:
        File content as text.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _load_json_memos(file_path: str) -> List[Dict[str, Any]]:
    """
    Load and flatten JSON memos into text blocks.

    Args:
        file_path: Path to the JSON file containing memos.

    Returns:
        List of dictionaries with flattened text and metadata.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Handle both single memo and list of memos
    if isinstance(data, dict):
        data = [data]

    memos = []
    for memo in data:
        # Flatten relevant fields into text
        text_parts = []

        if "facts" in memo:
            text_parts.append(f"Facts: {memo['facts']}")
        if "decision" in memo:
            text_parts.append(f"Decision: {memo['decision']}")
        if "cited_clauses" in memo:
            text_parts.append(f"Cited Clauses: {', '.join(memo['cited_clauses'])}")
        if "cited_sections" in memo:
            text_parts.append(f"Cited Sections: {', '.join(memo['cited_sections'])}")
        if "cited_exclusions" in memo:
            text_parts.append(f"Cited Exclusions: {', '.join(memo['cited_exclusions'])}")
        if "cited_regulations" in memo:
            text_parts.append(f"Cited Regulations: {', '.join(memo['cited_regulations'])}")

        # Include any other fields
        for key, value in memo.items():
            if key not in ["facts", "decision", "cited_clauses", "cited_sections",
                          "cited_exclusions", "cited_regulations"]:
                if isinstance(value, str):
                    text_parts.append(f"{key.replace('_', ' ').title()}: {value}")
                elif isinstance(value, (list, dict)):
                    text_parts.append(f"{key.replace('_', ' ').title()}: {json.dumps(value)}")

        memos.append({
            "text": "\n\n".join(text_parts),
            "raw_metadata": memo,
        })

    return memos


def _load_csv_memos(file_path: str) -> List[Dict[str, Any]]:
    """
    Load and flatten CSV memos into text blocks.
    Each row becomes a text block with key: value pairs.

    Args:
        file_path: Path to the CSV file containing memos.

    Returns:
        List of dictionaries with flattened text and metadata.
    """
    memos = []
    with open(file_path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            # Build a readable text block from all fields
            text_parts = []
            for key, value in row.items():
                if value and value.strip():
                    label = key.replace("_", " ").title()
                    text_parts.append(f"{label}: {value.strip()}")
            memos.append({
                "text": "\n".join(text_parts),
                "raw_metadata": dict(row),
            })
    return memos


def load_documents_from_manifest() -> List[Document]:
    """
    Load all documents from manifest sources.

    Returns:
        List of Document objects with text and metadata.
    """
    documents = []

    for source in iter_manifest_sources():
        file_path = source["full_path"]
        source_id = source["id"]
        doc_type = source["doc_type"]

        # Determine file type and load accordingly
        path_obj = Path(file_path)
        suffix = path_obj.suffix.lower()

        if suffix == ".pdf":
            text = _load_pdf(file_path)
            documents.append(Document(
                text=text,
                source_id=source_id,
                source_path=file_path,
                doc_type=doc_type,
                insurance_type=source["insurance_type"],
                product_code=source["product_code"],
                product_name=source["product_name"],
                claim_type=source["claim_type"],
                raw_metadata=source.get("metadata", {}),
            ))

        elif suffix == ".docx":
            text = _load_docx(file_path)
            documents.append(Document(
                text=text,
                source_id=source_id,
                source_path=file_path,
                doc_type=doc_type,
                insurance_type=source["insurance_type"],
                product_code=source["product_code"],
                product_name=source["product_name"],
                claim_type=source["claim_type"],
                raw_metadata=source.get("metadata", {}),
            ))

        elif suffix in [".md", ".markdown"]:
            text = _load_markdown(file_path)
            documents.append(Document(
                text=text,
                source_id=source_id,
                source_path=file_path,
                doc_type=doc_type,
                insurance_type=source["insurance_type"],
                product_code=source["product_code"],
                product_name=source["product_name"],
                claim_type=source["claim_type"],
                raw_metadata=source.get("metadata", {}),
            ))

        elif suffix == ".json":
            memos = _load_json_memos(file_path)
            for i, memo in enumerate(memos):
                documents.append(Document(
                    text=memo["text"],
                    source_id=f"{source_id}_{i}" if len(memos) > 1 else source_id,
                    source_path=file_path,
                    doc_type=doc_type,
                    insurance_type=source["insurance_type"],
                    product_code=source["product_code"],
                    product_name=source["product_name"],
                    claim_type=source["claim_type"],
                    raw_metadata={**source.get("metadata", {}), **memo["raw_metadata"]},
                ))

        elif suffix == ".csv":
            memos = _load_csv_memos(file_path)
            for i, memo in enumerate(memos):
                documents.append(Document(
                    text=memo["text"],
                    source_id=f"{source_id}_{i}" if len(memos) > 1 else source_id,
                    source_path=file_path,
                    doc_type=doc_type,
                    insurance_type=source["insurance_type"],
                    product_code=source["product_code"],
                    product_name=source["product_name"],
                    claim_type=source["claim_type"],
                    raw_metadata={**source.get("metadata", {}), **memo["raw_metadata"]},
                ))

    return documents
